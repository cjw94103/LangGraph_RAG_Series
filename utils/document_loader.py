import pandas as pd
import os

from docx import Document as Docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.table import Table as DocxTable
from docx.oxml.ns import qn

from langchain.schema import Document
from zipfile import ZipFile
from lxml import etree
from langchain_community.document_loaders.base import BaseLoader

from dataclasses import dataclass
from typing import List, Tuple, Iterable, Optional, Union, Set
import hashlib

from langchain.schema import Document

############################################ docx format ############################################

def _iter_blocks(doc: Docx): 
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, doc)

def _paragraph_split_by_pagebreak(p: DocxParagraph): 
    BR = qn('w:br')
    TYPE = qn('w:type')
    LR = qn('w:lastRenderedPageBreak')

    parts = []
    buf = []

    for r in p._p.r_lst:  # run-level 접근
        # run 내 lastRenderedPageBreak 탐지
        for lr in r.iterdescendants():
            if lr.tag == LR:
                # 경계 전까지 텍스트 flush
                text_before = ''.join(x.text or '' for x in r.iterchildren() if getattr(x, 'text', None))
                if text_before:
                    buf.append(text_before)
                parts.append(''.join(buf).strip())
                buf = []
                # lr 뒤 텍스트는 다음 루프에서 누적
                # 계속 진행
        # run 내 page break(<w:br w:type="page"/>) 처리
        br_elems = [e for e in r.findall('.//%s' % BR) if e.get(TYPE) == 'page']
        if br_elems:
            # br 앞 텍스트 flush
            text_before = ''.join(x.text or '' for x in r.iterchildren() if getattr(x, 'text', None))
            if text_before:
                buf.append(text_before)
            parts.append(''.join(buf).strip())
            buf = []
            continue

        # 일반 텍스트 누적
        if r.text:
            buf.append(r.text)

    if buf or not parts:
        parts.append(''.join(buf).strip())

    # 공백 페이지 방지
    parts = [s for s in parts if s is not None]
    return parts

def _table_to_markdown(tbl: DocxTable) -> str:
    rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
    if not rows:
        return ''
    header = rows[0] if any(x.strip() for x in rows[0]) else [f'col_{i+1}' for i in range(len(rows[0]))]
    df = pd.DataFrame(rows[1:], columns=header)
    return df.to_markdown(index=False)

def process_middle_docx(filepath: str):  
    d = Docx(filepath)

    pages = []          # 각 페이지는 [블록(str)] 리스트
    current = []        # 현 페이지 블록들

    for block in _iter_blocks(d):
        if isinstance(block, DocxParagraph):
            parts = _paragraph_split_by_pagebreak(block)
            for i, seg in enumerate(parts):
                if seg.strip():
                    current.append(seg.strip())
                # 분할 경계가 있고 뒤에 페이지가 이어지면 페이지 종료
                if i < len(parts) - 1:
                    pages.append(current)
                    current = []
        else:
            # 표는 원자적. 현재 페이지에 그대로 추가
            md = _table_to_markdown(block)
            if md.strip():
                current.append(md)

    # 마지막 페이지 flush
    if current:
        pages.append(current)

    # LangChain Document로 변환. 페이지 인덱스 1부터 부여.
    docs = []
    for i, blocks in enumerate(pages, start=1):
        # 텍스트와 표를 구분하기 위해 블록 사이에 빈 줄 1개로만 연결
        page_content = '\n\n'.join(blocks)
        docs.append(
            Document(
                page_content=page_content,
                metadata={"source": filepath, "page": i}
            )
        )
    return docs

############################################ hwpx format ############################################

def _ns_clean_xpath(root, xpath):  
    return root.xpath(xpath)

def _text_of(elem) -> str:  
    return ("".join(elem.itertext())).strip()

def _hash_md(s: str) -> str:  
    return hashlib.sha1(s.encode("utf-8", errors="ignore")).hexdigest()

@dataclass
class Block:   
    kind: str              
    xml: etree._Element    
    text: Optional[str]    
    md: Optional[str]      
    page_hint_y: Optional[float]  
    hard_break: bool       

def _read_pagedef(header_xml: Optional[etree._Element]) -> Tuple[Optional[float], Optional[float]]:  
    if header_xml is None:
        return None, None
    # pageDef 유사 태그 탐색
    nodes = _ns_clean_xpath(header_xml, ".//*[local-name()='pageDef' or local-name()='page-Def' or local-name()='pagedef']")
    if not nodes:
        return None, None
    node = nodes[0]
    # 단위가 mm, twip 등 다양할 수 있어 숫자만 취득
    def _num(attr):
        v = node.get(attr)
        if v is None: return None
        try:
            return float("".join(ch for ch in v if (ch.isdigit() or ch in ".-")))
        except:
            return None
    page_h = _num("height")
    top = _num("topMargin")
    bot = _num("bottomMargin")
    margin_sum = (top or 0.0) + (bot or 0.0)
    return page_h, margin_sum if (top is not None or bot is not None) else None

def _y_of_paragraph(p: etree._Element) -> Optional[float]:
    cand = _ns_clean_xpath(p, ".//*[local-name()='lineseg' or local-name()='lineSeg'][1]")
    if not cand:
        return None
    seg = cand[0]
    for key in ("y", "vertpos", "vertPos"):
        v = seg.get(key)
        if v is not None:
            try:
                return float("".join(ch for ch in v if (ch.isdigit() or ch in ".-")))
            except:
                continue
    return None

def _is_explicit_page_break(p: etree._Element) -> bool:  
    if _ns_clean_xpath(p, ".//*[local-name()='pageBreak' or local-name()='PageBreak']"):
        return True
    # run/br 유사 표기
    if _ns_clean_xpath(p, ".//*[local-name()='br' and (translate(@type,'PAGE','page')='page')]"):
        return True
    return False

def _tbl_to_md(tbl: etree._Element) -> str:   
    rows=[]
    for tr in _ns_clean_xpath(tbl, ".//*[local-name()='tr']"):
        cells=[_text_of(td) for td in _ns_clean_xpath(tr, ".//*[local-name()='tc']")]
        if cells: rows.append(cells)
    if not rows: 
        return ""
    # 헤더 생성
    head = rows[0]
    md = ["| " + " | ".join(head) + " |",
          "| " + " | ".join(["---"]*len(head)) + " |"]
    for r in rows[1:]:
        md.append("| " + " | ".join(r) + " |")
    return "\n".join(md)

def _iter_blocks_in_order(root: etree._Element) -> Iterable[Block]:  
    # 섹션 본문 루트 후보: hs:sec, hp:section 등
    # 직접 자식부터 DFS. 테이블 만나면 그 하위 p는 무시.
    def dfs(elem, inside_tbl=False):
        ln = etree.QName(elem).localname.lower()
        if ln == 'tbl':
            # 표는 원자 블록
            md = _tbl_to_md(elem)
            if md.strip():
                yield Block(kind='tbl', xml=elem, text=None, md=md,
                            page_hint_y=None, hard_break=False)
            # 표 내부는 별도 블록으로 내보내지 않음(중복 제거)
            return
        if ln == 'p' and not inside_tbl:
            t = _text_of(elem)
            if t:
                yield Block(kind='p', xml=elem, text=t, md=None,
                            page_hint_y=_y_of_paragraph(elem),
                            hard_break=_is_explicit_page_break(elem))
        # 자식 순회
        for ch in elem:
            ch_ln = etree.QName(ch).localname.lower()
            ch_inside_tbl = inside_tbl or (ln == 'tbl')
            # 표면 이미 처리했으므로 표 자식은 스킵
            if ln == 'tbl':
                continue
            yield from dfs(ch, inside_tbl=ch_inside_tbl)

    # 루트에서 시작
    for b in dfs(root, inside_tbl=False):
        yield b

def _paginate(blocks: List[Block], page_h: Optional[float], margin_sum: Optional[float]) -> List[List[Block]]: 
    pages: List[List[Block]] = []
    cur: List[Block] = []
    prev_y: Optional[float] = None
    acc_y: float = 0.0
    cutoff = (page_h - (margin_sum or 0.0)) if page_h else None

    for b in blocks:
        # 분기 조건 평가
        split = False
        if b.kind == 'p' and b.hard_break:
            split = True
        else:
            y = b.page_hint_y
            if y is not None:
                # 좌표가 작아지면 페이지 리셋으로 판단
                if prev_y is not None and y < prev_y:
                    split = True
                prev_y = y
                # 누적 기준(옵션)
                if cutoff:
                    acc_y += y if not split else 0.0
                    if acc_y > cutoff:
                        split = True
                        acc_y = 0.0
        if split and cur:
            pages.append(cur)
            cur = []
        cur.append(b)

    if cur:
        pages.append(cur)
    return pages

class HWPXLoaderPaged:
    def __init__(self, path: str):    
        self.path = path

    def _open_xml(self, z: ZipFile, name: str) -> etree._Element:   
        with z.open(name) as f:
            return etree.fromstring(f.read())

    def _iter_section_roots(self) -> Iterable[Tuple[str, etree._Element]]:       
        with ZipFile(self.path) as z:
            header_xml = None
            try:
                header_xml = self._open_xml(z, "Contents/header.xml")
            except KeyError:
                header_xml = None
            page_h, margin_sum = _read_pagedef(header_xml)

            # 섹션 파일들: content.hpf를 읽어 순서를 알 수도 있으나,
            # 관례적으로 section0.xml부터 존재. 정렬로 근사.
            section_names = sorted([n for n in z.namelist()
                                    if n.lower().startswith("contents/") and n.lower().endswith(".xml")
                                    and ("section" in n.lower())])
            for name in section_names:
                root = self._open_xml(z, name)
                yield name, root, page_h, margin_sum

    def load(self) -> List[Document]:   
        docs: List[Document] = []
        seen_tbl_hashes: Set[str] = set()

        for part, root, page_h, margin_sum in self._iter_section_roots():
            # 블록(순서 보존, 중복 제거)
            raw_blocks = list(_iter_blocks_in_order(root))

            # 표 중복 제거: 동일 MD 반복 방지
            blocks: List[Block] = []
            for b in raw_blocks:
                if b.kind == 'tbl':
                    h = _hash_md(b.md or "")
                    if h in seen_tbl_hashes:
                        continue
                    seen_tbl_hashes.add(h)
                blocks.append(b)

            # 페이지 단위 분할
            pages = _paginate(blocks, page_h, margin_sum)

            # 페이지 컨텐츠 직렬화: 블록 간 빈 줄 1개. 표는 MD 그대로.
            for page_idx, page_blocks in enumerate(pages, start=1):
                parts = []
                for b in page_blocks:
                    parts.append(b.text if b.kind == 'p' else b.md)
                content = "\n\n".join(x for x in parts if x)
                if not content.strip():
                    continue
                docs.append(
                    Document(
                        page_content=content,
                        metadata={
                            "source": self.path,
                            "part": part,
                            "page": page_idx,
                            "type": "page"
                        }
                    )
                )
        return docs

def process_fast_hwpx(filepath: str):  
    return HWPXLoaderPaged(filepath).load()

############################################ pdf format ############################################

from typing import List, Tuple
import pdfplumber
from langchain.schema import Document


def _calculate_overlap_area(box1: Tuple[float, float, float, float], 
                            box2: Tuple[float, float, float, float]) -> float:
    x_left = max(box1[0], box2[0])
    y_top = max(box1[1], box2[1])
    x_right = min(box1[2], box2[2])
    y_bottom = min(box1[3], box2[3])
    
    if x_right < x_left or y_bottom < y_top:
        return 0.0
    
    return (x_right - x_left) * (y_bottom - y_top)


def _overlaps_any(bbox: Tuple[float, float, float, float], 
                   table_bboxes: List[Tuple[float, float, float, float]], 
                   threshold: float = 0.5) -> bool:
    if not table_bboxes:
        return False
    
    text_area = (bbox[2] - bbox[0]) * (bbox[3] - bbox[1])
    if text_area == 0:
        return False
    
    for table_bbox in table_bboxes:
        overlap = _calculate_overlap_area(bbox, table_bbox)
        if overlap / text_area > threshold:
            return True
    
    return False


def load_pdf_documents(pdf_path: str) -> List[Document]:
    docs: List[Document] = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for pno, page in enumerate(pdf.pages, start=1):
            elements: List[Tuple[float, float, str, str]] = []  # (y0, x0, kind, content)
            table_bboxes: List[Tuple[float, float, float, float]] = []
            
            # 1) 표 탐지 및 Markdown 변환
            tables = page.find_tables()
            for table in tables:
                bbox = table.bbox  # (x0, y0, x1, y1)
                table_bboxes.append(bbox)
                
                # 표를 추출하여 Markdown 형식으로 변환
                extracted_table = table.extract()
                if extracted_table:
                    md = _convert_table_to_markdown(extracted_table)
                    if md:
                        y0, x0 = bbox[1], bbox[0]  # y0, x0 순서로 저장
                        elements.append((y0, x0, "table", md))
            
            # 2) 텍스트 추출 (단어 단위로 추출 후 라인으로 그룹화)
            words = page.extract_words(x_tolerance=3, y_tolerance=3)
            
            # 단어들을 y 좌표 기준으로 라인으로 그룹화
            lines = _group_words_into_lines(words, y_tolerance=3)
            
            for line_y, line_x0, line_text, line_bbox in lines:
                # 표 영역과 겹치지 않는 텍스트만 추가
                if not _overlaps_any(line_bbox, table_bboxes):
                    elements.append((line_y, line_x0, "text", line_text))
            
            # 3) 위→아래(동일 y에서는 x) 정렬 후 병합
            if not elements:
                continue
                
            elements.sort(key=lambda e: (e[0], e[1]))
            
            parts: List[str] = []
            table_counter = 0
            prev_kind = None
            
            for _, _, kind, content in elements:
                if kind == "table":
                    table_counter += 1
                    parts.append(f"\n{content}")
                    prev_kind = "table"
                else:  # text
                    if prev_kind == "text":
                        parts.append(content)
                    else:
                        parts.append("\n" + content)
                    prev_kind = "text"
            
            page_content = "\n\n".join(parts)
            
            docs.append(Document(
                page_content=page_content,
                metadata={
                    "source": pdf_path,
                    "page": pno,
                    "engine": "pdfplumber",
                    "has_tables": bool(tables),
                    "num_tables": len(tables),
                    "order": "top-to-bottom by (y,x)",
                }
            ))
    
    return docs


def _group_words_into_lines(words: List[dict], y_tolerance: float = 3) -> List[Tuple[float, float, str, Tuple[float, float, float, float]]]:
    if not words:
        return []
    
    # y 좌표 기준으로 정렬
    sorted_words = sorted(words, key=lambda w: (w['top'], w['x0']))
    
    lines = []
    current_line_words = [sorted_words[0]]
    current_y = sorted_words[0]['top']
    
    for word in sorted_words[1:]:
        # 같은 라인인지 판단 (y 좌표 차이가 tolerance 이내)
        if abs(word['top'] - current_y) <= y_tolerance:
            current_line_words.append(word)
        else:
            # 현재 라인 완성
            lines.append(_create_line_from_words(current_line_words))
            # 새 라인 시작
            current_line_words = [word]
            current_y = word['top']
    
    # 마지막 라인 추가
    if current_line_words:
        lines.append(_create_line_from_words(current_line_words))
    
    return lines


def _create_line_from_words(words: List[dict]) -> Tuple[float, float, str, Tuple[float, float, float, float]]:
    # x 좌표 기준으로 정렬
    words = sorted(words, key=lambda w: w['x0'])
    
    text = ' '.join(w['text'] for w in words)
    y0 = min(w['top'] for w in words)
    x0 = min(w['x0'] for w in words)
    x1 = max(w['x1'] for w in words)
    y1 = max(w['bottom'] for w in words)
    
    bbox = (x0, y0, x1, y1)
    
    return (y0, x0, text, bbox)


def _convert_table_to_markdown(table_data: List[List]) -> str:
    if not table_data:
        return ""
    
    # None 값을 빈 문자열로 변환
    clean_data = []
    for row in table_data:
        clean_row = [str(cell) if cell is not None else "" for cell in row]
        clean_data.append(clean_row)
    
    if not clean_data:
        return ""
    
    # Markdown 표 생성
    lines = []
    
    # 헤더 행
    header = "| " + " | ".join(clean_data[0]) + " |"
    lines.append(header)
    
    # 구분선
    separator = "| " + " | ".join(["---"] * len(clean_data[0])) + " |"
    lines.append(separator)
    
    # 데이터 행
    for row in clean_data[1:]:
        # 열 개수가 헤더와 다를 경우 조정
        if len(row) < len(clean_data[0]):
            row = row + [""] * (len(clean_data[0]) - len(row))
        elif len(row) > len(clean_data[0]):
            row = row[:len(clean_data[0])]
        
        data_row = "| " + " | ".join(row) + " |"
        lines.append(data_row)
    
    return "\n".join(lines)

############################################ txt format ############################################

def load_txt_as_documents(path: str, encoding: str = "utf-8") -> List[Document]:
    
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    
    with open(path, "r", encoding=encoding) as f:
        text = f.read().strip()

    if not text:
        return []

    doc = Document(
        page_content=text,
        metadata={"source": path, "type": "txt"}
    )
    return [doc]

############################################ md format ############################################

def load_md_as_documents(path: str, encoding: str = "utf-8") -> List[Document]:
    
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    
    with open(path, "r", encoding=encoding) as f:
        text = f.read().strip()

    if not text:
        return []

    doc = Document(
        page_content=text,
        metadata={"source": path, "type": "markdown"}
    )
    return [doc]